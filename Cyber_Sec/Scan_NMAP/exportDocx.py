from collections import Counter
from docx import Document
from docx.shared import Inches

def generate_docx_from_log(log_filename, docx_filename, successful_logins, discovered_services):
    with open(log_filename) as f:
        log_lines = f.read().splitlines()

    # Extract ports
    ports = []
    for line in log_lines:
        if line.startswith('Discovered '):
            try:
                _, service, _, host_port = line.split()
                host, port = host_port.split(':')
                ports.append((host, port))
            except ValueError:
                continue

    # Count ports
    port_counter = Counter(port for ip, port in ports)

    # Create Word document
    doc = Document()

    # Add table for port statistics
    doc.add_heading('Thống kê số lượng port được mở', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Số thứ tự'
    hdr_cells[1].text = 'Port'
    hdr_cells[2].text = 'Số lượng'
    for i, (port, count) in enumerate(port_counter.items(), start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = port
        row_cells[2].text = str(count)

    # Add table for successful logins
    doc.add_heading('Thống kê số lượng Port nguy hiểm được mở', level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Số thứ tự'
    hdr_cells[1].text = 'IP address'
    hdr_cells[2].text = 'Port'
    hdr_cells[3].text = 'Username'
    hdr_cells[4].text = 'Password'
    for i, (ip, port, username, password) in enumerate(successful_logins, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = ip
        row_cells[2].text = str(port)
        row_cells[3].text = username if username != 'anonymous' else 'anonymous (no password needed)'
        row_cells[4].text = password if username != 'anonymous' else ''
    
    doc.add_heading('Các máy tính mở', level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'IP'
    hdr_cells[1].text = 'Port'
    for host, service, port in discovered_services:
        row_cells = table.add_row().cells
        row_cells[0].text = host
        row_cells[1].text = str(port)
    
    # Save the document
    doc.save(docx_filename)

